# backend/scripts/generar_anexo.py - GENERADOR DE ANEXOS DOCX CON TABLA DE EQUIPOS
import sys
import json
import os
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def reemplazar_texto(doc, marcador, nuevo_texto):
    for p in doc.paragraphs:
        if marcador in p.text:
            texto = p.text.replace(marcador, nuevo_texto)
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = texto
            else:
                p.add_run(texto)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_texto(cell, marcador, nuevo_texto)

def reemplazar_equipos_con_tabla(doc, equipos_list):
    target_p = None
    for p in doc.paragraphs:
        if "{{equipos}}" in p.text:
            target_p = p
            break
            
    if not target_p:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "{{equipos}}" in p.text:
                            target_p = p
                            break

    if not target_p:
        return

    if not equipos_list or len(equipos_list) == 0:
        target_p.text = "Sin equipos asignados."
        return

    table = doc.add_table(rows=1, cols=6)
    set_table_borders(table)
    
    hdr_cells = table.rows[0].cells
    headers = ["Tipo", "Marca", "Modelo", "N° Serie", "Estado", "Observaciones"]
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "2A3284")
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(9.5)
                r.font.name = "Calibri"

    for idx, eq in enumerate(equipos_list):
        row_cells = table.add_row().cells
        bg_color = "F8FAFC" if idx % 2 == 1 else "FFFFFF"
        
        tipo_val = str(eq.get("tipo") or eq.get("nombre") or "Equipo")
        marca_val = str(eq.get("marca") or "N/A")
        modelo_val = str(eq.get("modelo") or "N/A")
        serie_val = str(eq.get("numero_serie") or eq.get("serie") or "N/A")
        estado_val = str(eq.get("estado") or eq.get("condicion") or "BUENO")
        obs_val = str(eq.get("observaciones") or eq.get("comentario") or "")
        
        row_data = [tipo_val, marca_val, modelo_val, serie_val, estado_val, obs_val]
        
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            set_cell_background(row_cells[i], bg_color)
            for p in row_cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in [0, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Calibri"
                    r.font.color.rgb = RGBColor(30, 41, 59)

    target_p._p.addprevious(table._tbl)
    target_p._p.getparent().remove(target_p._p)

def main():
    if len(sys.argv) < 2:
        print(json.dumps({"success": False, "error": "Faltan argumentos JSON"}))
        sys.exit(1)

    try:
        raw_json = sys.argv[1]
        data = json.loads(raw_json)
        
        template_path = data["template_path"]
        output_path = data["output_path"]
        fecha = data.get("fecha", "")
        nombre = data.get("nombre", "")
        rut = data.get("rut", "")
        equipos_list = data.get("equipos_list", [])

        if not os.path.exists(template_path):
            print(json.dumps({"success": False, "error": f"Plantilla no encontrada: {template_path}"}))
            sys.exit(1)

        doc = docx.Document(template_path)
        reemplazar_texto(doc, "{{fecha}}", fecha)
        reemplazar_texto(doc, "{{nombre}}", nombre)
        reemplazar_texto(doc, "{{rut}}", rut)
        
        reemplazar_equipos_con_tabla(doc, equipos_list)

        out_dir = os.path.dirname(output_path)
        if out_dir and not os.path.exists(out_dir):
            os.makedirs(out_dir, exist_ok=True)

        doc.save(output_path)
        print(json.dumps({"success": True, "output_path": output_path}))
    except Exception as e:
        print(json.dumps({"success": False, "error": str(e)}))
        sys.exit(1)

if __name__ == "__main__":
    main()
